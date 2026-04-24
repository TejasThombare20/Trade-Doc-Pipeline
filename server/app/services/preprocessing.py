"""Document preprocessing service.

Given raw bytes from storage, classify the file and produce a normalized
representation the extractor can feed to a vision LLM:

  * text         — whatever text we could pull natively
  * images       — base64-encoded page images for vision inference
  * page_count   — for logging / UI
  * source_kind  — "pdf_text" | "pdf_scanned" | "image" | "docx"
"""

from __future__ import annotations

import asyncio
import base64
import io
from dataclasses import dataclass, field
from typing import Literal

import pdfplumber
from pdf2image import convert_from_bytes
from PIL import Image

from app.core.errors import PreprocessingError
from app.core.logging import get_logger

logger = get_logger(__name__)

SourceKind = Literal["pdf_text", "pdf_scanned", "image", "docx"]

_MIN_CHARS_PER_PAGE_FOR_TEXT_PDF = 80
_MAX_VISION_PAGES = 4
_RENDER_DPI = 200


@dataclass
class PreprocessedDocument:
    source_kind: SourceKind
    mime_type: str
    page_count: int
    text: str = ""
    images_b64: list[str] = field(default_factory=list)
    notes: str | None = None


class PreprocessingService:
    """Stateless service; instantiate once and reuse."""

    async def preprocess(
        self,
        data: bytes,
        mime_type: str,
        original_name: str,
    ) -> PreprocessedDocument:
        """Dispatch on mime type. Heavy CPU work runs in a thread."""
        mime = mime_type.lower()
        if mime == "application/pdf" or original_name.lower().endswith(".pdf"):
            return await asyncio.to_thread(_preprocess_pdf, data)
        if mime.startswith("image/"):
            return await asyncio.to_thread(_preprocess_image, data, mime)
        if mime in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ) or original_name.lower().endswith(".docx"):
            return await asyncio.to_thread(_preprocess_docx, data)

        raise PreprocessingError(
            f"unsupported file type: mime={mime_type} name={original_name}"
        )


# Module-level singleton + alias for backward-compat
_preprocessing_service: PreprocessingService = PreprocessingService()


def get_preprocessing_service() -> PreprocessingService:
    return _preprocessing_service


async def preprocess(data: bytes, mime_type: str, original_name: str) -> PreprocessedDocument:
    """Backward-compat module-level function."""
    return await _preprocessing_service.preprocess(data, mime_type, original_name)


# ─── pure helpers ────────────────────────────────────────────────────────────

def _preprocess_pdf(data: bytes) -> PreprocessedDocument:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            page_count = len(pdf.pages)
            parts: list[str] = []
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                parts.append(extracted)
            text = "\n\n".join(parts).strip()
    except Exception as exc:
        raise PreprocessingError(f"pdf_parse_failed: {exc}") from exc

    if page_count == 0:
        raise PreprocessingError("pdf has zero pages")

    char_density = len(text.replace(" ", "").replace("\n", "")) / page_count
    if char_density >= _MIN_CHARS_PER_PAGE_FOR_TEXT_PDF:
        images_b64 = _render_pdf_pages(data, max_pages=1)
        return PreprocessedDocument(
            source_kind="pdf_text",
            mime_type="application/pdf",
            page_count=page_count,
            text=text,
            images_b64=images_b64,
        )

    images_b64 = _render_pdf_pages(data, max_pages=_MAX_VISION_PAGES)
    return PreprocessedDocument(
        source_kind="pdf_scanned",
        mime_type="application/pdf",
        page_count=page_count,
        text=text,
        images_b64=images_b64,
        notes=(
            f"Low text density ({char_density:.0f} chars/page); "
            "treating as scanned and sending pages to vision."
        ),
    )


def _render_pdf_pages(data: bytes, *, max_pages: int) -> list[str]:
    try:
        pages = convert_from_bytes(data, dpi=_RENDER_DPI, fmt="png")
    except Exception as exc:
        raise PreprocessingError(
            "pdf_render_failed (is poppler-utils installed?): " + str(exc)
        ) from exc
    pages = pages[:max_pages]
    return [_png_b64(img) for img in pages]


def _preprocess_image(data: bytes, mime: str) -> PreprocessedDocument:
    try:
        img = Image.open(io.BytesIO(data))
        img.load()
    except Exception as exc:
        raise PreprocessingError(f"image_parse_failed: {exc}") from exc

    return PreprocessedDocument(
        source_kind="image",
        mime_type=mime,
        page_count=1,
        text="",
        images_b64=[_png_b64(img)],
    )


def _preprocess_docx(data: bytes) -> PreprocessedDocument:
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
    except Exception as exc:
        raise PreprocessingError(f"docx_parse_failed: {exc}") from exc

    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    text = "\n".join(parts).strip()
    if not text:
        raise PreprocessingError("docx had no readable text")

    return PreprocessedDocument(
        source_kind="docx",
        mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        page_count=1,
        text=text,
    )


def _png_b64(img: Image.Image) -> str:
    buf = io.BytesIO()
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    img.save(buf, format="PNG", optimize=True)
    return base64.b64encode(buf.getvalue()).decode("ascii")
